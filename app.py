import os
import io
import json
import re
import tempfile
from pathlib import Path

from flask import Flask, request, jsonify, render_template, Response, stream_with_context
from dotenv import load_dotenv
import pdfplumber
import anthropic

load_dotenv()

app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = 50 * 1024 * 1024  # 50 MB upload limit

client = anthropic.Anthropic(api_key=os.environ.get("ANTHROPIC_API_KEY"))

MAX_CONTRACT_CHARS = 150_000

ANALYSIS_PROMPT = """You are an expert Australian conveyancing analyst. You have been given the text of a Contract of Sale for a residential property in Australia.

Your task is to analyse this contract thoroughly and return a structured JSON response. Do not provide legal advice — frame your analysis as things to be aware of and discuss with a licensed conveyancer.

Analyse the following contract text and return ONLY a valid JSON object with exactly these keys:

{{
  "summary": "A plain-English summary of the property, parties, purchase price, settlement date, and overall deal structure. 3-5 paragraphs.",
  "key_risks": [
    {{"risk": "Short title of risk", "detail": "Detailed explanation of the risk and why it matters"}}
  ],
  "special_conditions": [
    {{"condition": "Title or summary of special condition", "detail": "What it means in plain English and any concerns"}}
  ],
  "title_issues": {{
    "overview": "Summary of title, easements, covenants, caveats, and encumbrances",
    "items": [
      {{"type": "Easement|Covenant|Caveat|Restriction|Other", "description": "What it is and how it may affect the buyer"}}
    ]
  }},
  "owners_corporation": {{
    "applicable": true or false,
    "overview": "Summary of owners corporation situation, fees, and any known issues. Set to null if not applicable.",
    "issues": ["List of specific concerns if any"]
  }},
  "settlement_risks": [
    {{"risk": "Risk title", "detail": "Explanation"}}
  ],
  "finance_deposit": {{
    "deposit_amount": "Deposit amount and when due",
    "deposit_holder": "Who holds the deposit (vendor's solicitor, agent, etc.)",
    "finance_clause": "Details of any finance clause — date, amount, lender if specified",
    "cooling_off": "Cooling off period details (duration, penalty if applicable, any waiver)",
    "key_dates": [
      {{"date_type": "Label", "date": "Date or timeframe"}}
    ]
  }},
  "due_diligence": [
    {{"item": "Checklist item", "status": "Do before settlement|Confirm with conveyancer|Review carefully|Not applicable", "detail": "What to check and why"}}
  ],
  "conveyancer_questions": [
    {{"question": "Specific question to ask your conveyancer", "context": "Why this matters"}}
  ],
  "deal_memo": {{
    "headline": "One-sentence deal verdict",
    "situation": "What is being purchased, from whom, and at what price",
    "complication": "The key issues or risks in this contract",
    "resolution": "What needs to happen for this deal to proceed safely",
    "recommendation": "Overall recommendation: PROCEED WITH CAUTION | CLARIFY BEFORE PROCEEDING | HIGH RISK — SEEK ADVICE",
    "key_numbers": [
      {{"label": "Metric or number", "value": "Value"}}
    ],
    "next_steps": ["Ordered list of immediate actions to take"]
  }}
}}

Contract text to analyse:

{contract_text}

Return ONLY the JSON object. No preamble, no explanation, no markdown code blocks."""


def extract_pdf_text(file_bytes: bytes) -> str:
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        pages = []
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
    return "\n\n".join(pages)


def parse_analysis_json(raw: str) -> dict:
    # Strip markdown code fences if present
    raw = raw.strip()
    if raw.startswith("```"):
        raw = re.sub(r"^```[a-z]*\n?", "", raw)
        raw = re.sub(r"\n?```$", "", raw)

    # Find the outermost JSON object
    start = raw.find("{")
    end = raw.rfind("}") + 1
    if start == -1 or end == 0:
        raise ValueError("No JSON object found in response")
    return json.loads(raw[start:end])


@app.route("/")
def index():
    return render_template("index.html")


@app.route("/analyse", methods=["POST"])
def analyse():
    if "contract" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400

    file = request.files["contract"]
    if not file.filename or not file.filename.lower().endswith(".pdf"):
        return jsonify({"error": "Please upload a PDF file"}), 400

    file_bytes = file.read()
    if not file_bytes:
        return jsonify({"error": "Uploaded file is empty"}), 400

    try:
        contract_text = extract_pdf_text(file_bytes)
    except Exception as e:
        return jsonify({"error": f"Could not read PDF: {str(e)}. If the PDF is scanned/image-based, text extraction is not supported."}), 400

    if not contract_text.strip():
        return jsonify({"error": "No text could be extracted from this PDF. It may be a scanned document — please use a text-based PDF."}), 400

    if len(contract_text) > MAX_CONTRACT_CHARS:
        contract_text = contract_text[:MAX_CONTRACT_CHARS] + "\n\n[Contract truncated due to length — first 150,000 characters analysed]"

    prompt = ANALYSIS_PROMPT.format(contract_text=contract_text)

    def generate():
        full_response = []
        try:
            with client.messages.stream(
                model="claude-opus-4-8",
                max_tokens=8000,
                thinking={"type": "adaptive"},
                messages=[{"role": "user", "content": prompt}],
            ) as stream:
                for text in stream.text_stream:
                    full_response.append(text)
                    # Send heartbeat so browser knows we're alive
                    yield f"data: {json.dumps({'type': 'chunk', 'text': text})}\n\n"

            raw = "".join(full_response)
            try:
                analysis = parse_analysis_json(raw)
                yield f"data: {json.dumps({'type': 'complete', 'analysis': analysis})}\n\n"
            except Exception as e:
                yield f"data: {json.dumps({'type': 'error', 'error': f'Failed to parse analysis: {str(e)}'})}\n\n"

        except anthropic.APIStatusError as e:
            yield f"data: {json.dumps({'type': 'error', 'error': f'Claude API error: {e.message}'})}\n\n"
        except Exception as e:
            yield f"data: {json.dumps({'type': 'error', 'error': str(e)})}\n\n"

    return Response(stream_with_context(generate()), mimetype="text/event-stream",
                    headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"})


@app.route("/export/markdown", methods=["POST"])
def export_markdown():
    data = request.get_json()
    if not data or "analysis" not in data:
        return jsonify({"error": "No analysis data provided"}), 400

    a = data["analysis"]
    lines = []

    lines.append("# Conveyancer Co-pilot — Contract Analysis Report")
    lines.append("> **Disclaimer:** This analysis is not legal advice. Please discuss all findings with a licensed conveyancer before proceeding.")
    lines.append("")

    lines.append("## 1. Plain-English Summary")
    lines.append(a.get("summary", ""))
    lines.append("")

    lines.append("## 2. Key Risks & Red Flags")
    for item in a.get("key_risks", []):
        lines.append(f"### {item.get('risk', '')}")
        lines.append(item.get("detail", ""))
        lines.append("")

    lines.append("## 3. Special Conditions Review")
    for item in a.get("special_conditions", []):
        lines.append(f"### {item.get('condition', '')}")
        lines.append(item.get("detail", ""))
        lines.append("")

    lines.append("## 4. Title, Easements & Covenants")
    ti = a.get("title_issues", {})
    lines.append(ti.get("overview", ""))
    for item in ti.get("items", []):
        lines.append(f"- **{item.get('type', '')}:** {item.get('description', '')}")
    lines.append("")

    lines.append("## 5. Owners Corporation")
    oc = a.get("owners_corporation", {})
    if oc.get("applicable"):
        lines.append(oc.get("overview", ""))
        for issue in oc.get("issues", []):
            lines.append(f"- {issue}")
    else:
        lines.append("Not applicable to this property.")
    lines.append("")

    lines.append("## 6. Settlement & Default Risks")
    for item in a.get("settlement_risks", []):
        lines.append(f"### {item.get('risk', '')}")
        lines.append(item.get("detail", ""))
        lines.append("")

    lines.append("## 7. Finance & Deposit Obligations")
    fd = a.get("finance_deposit", {})
    lines.append(f"**Deposit:** {fd.get('deposit_amount', '')}")
    lines.append(f"**Deposit Holder:** {fd.get('deposit_holder', '')}")
    lines.append(f"**Finance Clause:** {fd.get('finance_clause', '')}")
    lines.append(f"**Cooling Off:** {fd.get('cooling_off', '')}")
    lines.append("")
    lines.append("### Key Dates")
    for kd in fd.get("key_dates", []):
        lines.append(f"- **{kd.get('date_type', '')}:** {kd.get('date', '')}")
    lines.append("")

    lines.append("## 8. Due Diligence Checklist")
    for item in a.get("due_diligence", []):
        lines.append(f"- [{item.get('status', '')}] **{item.get('item', '')}** — {item.get('detail', '')}")
    lines.append("")

    lines.append("## 9. Questions for Your Conveyancer")
    for i, item in enumerate(a.get("conveyancer_questions", []), 1):
        lines.append(f"{i}. **{item.get('question', '')}**")
        lines.append(f"   _{item.get('context', '')}_")
        lines.append("")

    lines.append("## 10. Deal Memo")
    dm = a.get("deal_memo", {})
    lines.append(f"**{dm.get('headline', '')}**")
    lines.append("")
    lines.append(f"**Situation:** {dm.get('situation', '')}")
    lines.append("")
    lines.append(f"**Complication:** {dm.get('complication', '')}")
    lines.append("")
    lines.append(f"**Resolution:** {dm.get('resolution', '')}")
    lines.append("")
    lines.append(f"**Recommendation:** `{dm.get('recommendation', '')}`")
    lines.append("")
    lines.append("### Key Numbers")
    for kn in dm.get("key_numbers", []):
        lines.append(f"| {kn.get('label', '')} | {kn.get('value', '')} |")
    lines.append("")
    lines.append("### Next Steps")
    for step in dm.get("next_steps", []):
        lines.append(f"1. {step}")
    lines.append("")

    md_content = "\n".join(lines)
    return Response(
        md_content,
        mimetype="text/markdown",
        headers={"Content-Disposition": "attachment; filename=contract-analysis.md"}
    )


@app.route("/export/docx", methods=["POST"])
def export_docx():
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    data = request.get_json()
    if not data or "analysis" not in data:
        return jsonify({"error": "No analysis data provided"}), 400

    a = data["analysis"]
    doc = Document()

    # Title
    title = doc.add_heading("Conveyancer Co-pilot — Contract Analysis", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    disclaimer = doc.add_paragraph("Disclaimer: This analysis is not legal advice. Please discuss all findings with a licensed conveyancer before proceeding.")
    disclaimer.runs[0].italic = True
    disclaimer.runs[0].font.color.rgb = RGBColor(0x88, 0x44, 0x00)

    doc.add_paragraph("")

    def add_section(num, title_text, content_fn):
        doc.add_heading(f"{num}. {title_text}", level=1)
        content_fn()
        doc.add_paragraph("")

    def section_summary():
        doc.add_paragraph(a.get("summary", ""))

    def section_risks():
        for item in a.get("key_risks", []):
            p = doc.add_paragraph()
            run = p.add_run(item.get("risk", ""))
            run.bold = True
            doc.add_paragraph(item.get("detail", ""))

    def section_special():
        for item in a.get("special_conditions", []):
            p = doc.add_paragraph()
            run = p.add_run(item.get("condition", ""))
            run.bold = True
            doc.add_paragraph(item.get("detail", ""))

    def section_title():
        ti = a.get("title_issues", {})
        doc.add_paragraph(ti.get("overview", ""))
        for item in ti.get("items", []):
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(f"{item.get('type', '')}: ")
            run.bold = True
            p.add_run(item.get("description", ""))

    def section_oc():
        oc = a.get("owners_corporation", {})
        if oc.get("applicable"):
            doc.add_paragraph(oc.get("overview", "") or "")
            for issue in oc.get("issues", []):
                doc.add_paragraph(issue, style="List Bullet")
        else:
            doc.add_paragraph("Not applicable to this property.")

    def section_settlement():
        for item in a.get("settlement_risks", []):
            p = doc.add_paragraph()
            run = p.add_run(item.get("risk", ""))
            run.bold = True
            doc.add_paragraph(item.get("detail", ""))

    def section_finance():
        fd = a.get("finance_deposit", {})
        for label, val in [
            ("Deposit", fd.get("deposit_amount", "")),
            ("Deposit Holder", fd.get("deposit_holder", "")),
            ("Finance Clause", fd.get("finance_clause", "")),
            ("Cooling Off", fd.get("cooling_off", "")),
        ]:
            p = doc.add_paragraph()
            p.add_run(f"{label}: ").bold = True
            p.add_run(val or "")
        doc.add_paragraph("")
        doc.add_heading("Key Dates", level=2)
        for kd in fd.get("key_dates", []):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"{kd.get('date_type', '')}: ").bold = True
            p.add_run(kd.get("date", ""))

    def section_due_diligence():
        for item in a.get("due_diligence", []):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"[{item.get('status', '')}] ").bold = True
            p.add_run(f"{item.get('item', '')} — {item.get('detail', '')}")

    def section_questions():
        for i, item in enumerate(a.get("conveyancer_questions", []), 1):
            p = doc.add_paragraph()
            p.add_run(f"{i}. {item.get('question', '')}").bold = True
            doc.add_paragraph(item.get("context", ""))

    def section_deal_memo():
        dm = a.get("deal_memo", {})
        p = doc.add_paragraph()
        p.add_run(dm.get("headline", "")).bold = True

        for label, key in [
            ("Situation", "situation"),
            ("Complication", "complication"),
            ("Resolution", "resolution"),
            ("Recommendation", "recommendation"),
        ]:
            p2 = doc.add_paragraph()
            p2.add_run(f"{label}: ").bold = True
            p2.add_run(dm.get(key, ""))

        doc.add_heading("Key Numbers", level=2)
        for kn in dm.get("key_numbers", []):
            p3 = doc.add_paragraph(style="List Bullet")
            p3.add_run(f"{kn.get('label', '')}: ").bold = True
            p3.add_run(kn.get("value", ""))

        doc.add_heading("Next Steps", level=2)
        for step in dm.get("next_steps", []):
            doc.add_paragraph(step, style="List Number")

    add_section(1, "Plain-English Summary", section_summary)
    add_section(2, "Key Risks & Red Flags", section_risks)
    add_section(3, "Special Conditions Review", section_special)
    add_section(4, "Title, Easements & Covenants", section_title)
    add_section(5, "Owners Corporation", section_oc)
    add_section(6, "Settlement & Default Risks", section_settlement)
    add_section(7, "Finance & Deposit Obligations", section_finance)
    add_section(8, "Due Diligence Checklist", section_due_diligence)
    add_section(9, "Questions for Your Conveyancer", section_questions)
    add_section(10, "Deal Memo", section_deal_memo)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    return Response(
        buf.read(),
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=contract-analysis.docx"}
    )


if __name__ == "__main__":
    api_key = os.environ.get("ANTHROPIC_API_KEY")
    if not api_key:
        print("ERROR: ANTHROPIC_API_KEY not set. Copy .env.example to .env and add your key.")
    else:
        print("Starting Conveyancer Co-pilot at http://localhost:5000")
    app.run(debug=False, port=5000)
